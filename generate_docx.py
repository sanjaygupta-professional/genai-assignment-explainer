#!/usr/bin/env python3.12
"""Generate a Word document from the Assignment 1 workbook HTML."""

import re
import io
from pathlib import Path
from html.parser import HTMLParser
from datetime import date

import cairosvg
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT


BASE = Path("/home/opc/genai-assignment-explainer")
HTML_FILE = BASE / "assignment1-workbook.html"
OUTPUT = BASE / "outputs" / "Assignment1_DBA_GenAI_Professor.docx"
OUTPUT.parent.mkdir(parents=True, exist_ok=True)

# Image paths
DIAGRAM_SVG = BASE / "images" / "diagrams" / "rag-pipeline-architecture.svg"
SCREENSHOT_DIR = BASE / "images" / "screenshots"
SCREENSHOTS = {
    "02-test-a-attention.png": "Test A: Attention Mechanism Query — detailed explanation with 8 cited sources",
    "03-test-b-outofscope.png": "Test B: Out-of-Scope Refusal — graceful handling of non-course query",
    "04-test-c-synthesis.png": "Test C: Multi-Document Synthesis — comparing BERT, GPT, and T5 architectures",
}


class HTMLTextExtractor(HTMLParser):
    """Simple HTML to plain text converter."""
    def __init__(self):
        super().__init__()
        self._text = []
        self._skip = False
        self._in_code = False

    def handle_starttag(self, tag, attrs):
        if tag in ("script", "style", "nav", "header"):
            self._skip = True
        if tag == "code":
            self._in_code = True
        if tag == "br":
            self._text.append("\n")

    def handle_endtag(self, tag):
        if tag in ("script", "style", "nav", "header"):
            self._skip = False
        if tag == "code":
            self._in_code = False
        if tag in ("p", "li", "tr", "h4", "h3", "div"):
            self._text.append("\n")

    def handle_data(self, data):
        if not self._skip:
            self._text.append(data)

    def handle_entityref(self, name):
        entities = {"mdash": "—", "ndash": "–", "bull": "•", "amp": "&",
                     "lt": "<", "gt": ">", "nbsp": " ", "hellip": "…"}
        self._text.append(entities.get(name, f"&{name};"))

    def handle_charref(self, name):
        try:
            if name.startswith("x"):
                c = chr(int(name[1:], 16))
            else:
                c = chr(int(name))
            self._text.append(c)
        except (ValueError, OverflowError):
            self._text.append(f"&#{name};")

    def get_text(self):
        return "".join(self._text)


def html_to_text(html_str):
    """Convert HTML fragment to plain text."""
    parser = HTMLTextExtractor()
    parser.feed(html_str)
    text = parser.get_text()
    # Clean up excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def extract_your_work_sections(html):
    """Extract all 'Your Work' div content blocks from the workbook HTML."""
    # Find all your-work-placeholder divs
    pattern = r'<div class="your-work-placeholder">(.*?)</div>\s*</div>'
    # Need DOTALL for multi-line matching
    matches = re.findall(pattern, html, re.DOTALL)
    return matches


def extract_section_titles(html):
    """Extract section numbers and titles."""
    pattern = r'<div class="section-num">(\d+)</div>\s*<div class="section-header-text">\s*<h2>(.*?)</h2>'
    matches = re.findall(pattern, html, re.DOTALL)
    import html as html_mod
    return [(num, html_mod.unescape(re.sub(r'<[^>]+>', '', title).strip())) for num, title in matches]


def extract_tables_from_html(html_fragment):
    """Extract table data from HTML fragment. Returns list of (headers, rows)."""
    tables = []
    table_pattern = r'<table[^>]*>(.*?)</table>'
    for table_match in re.finditer(table_pattern, html_fragment, re.DOTALL):
        table_html = table_match.group(1)

        headers = []
        header_pattern = r'<th[^>]*>(.*?)</th>'
        headers = [re.sub(r'<[^>]+>', '', h).strip()
                    for h in re.findall(header_pattern, table_html, re.DOTALL)]

        rows = []
        row_pattern = r'<tr>(.*?)</tr>'
        for row_match in re.findall(row_pattern, table_html, re.DOTALL):
            if '<th' in row_match:
                continue  # skip header row
            cells = [re.sub(r'<[^>]+>', '', c).strip()
                     for c in re.findall(r'<td[^>]*>(.*?)</td>', row_match, re.DOTALL)]
            if cells:
                rows.append(cells)

        if headers or rows:
            tables.append((headers, rows))

    return tables


def add_table_to_doc(doc, headers, rows):
    """Add a formatted table to the document."""
    if not headers and not rows:
        return
    ncols = len(headers) if headers else len(rows[0]) if rows else 0
    if ncols == 0:
        return

    nrows = (1 if headers else 0) + len(rows)
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    row_idx = 0
    if headers:
        for j, h in enumerate(headers):
            cell = table.cell(0, j)
            cell.text = h
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        row_idx = 1

    for row_data in rows:
        for j, val in enumerate(row_data):
            if j < ncols:
                cell = table.cell(row_idx, j)
                cell.text = val
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
        row_idx += 1

    doc.add_paragraph()  # spacing after table


def process_section_content(doc, html_fragment, section_num):
    """Process a Your Work HTML fragment into document paragraphs and tables."""
    # Split content into segments: text blocks and tables
    # First extract tables and their positions
    table_data = extract_tables_from_html(html_fragment)

    # Remove tables from HTML to get text portions
    text_html = re.sub(r'<table[^>]*>.*?</table>', '\n[TABLE]\n', html_fragment, flags=re.DOTALL)

    # Split by h4 headings
    import html as html_mod
    h4_pattern = r'<h4>(.*?)</h4>'
    parts = re.split(h4_pattern, text_html, flags=re.DOTALL)

    table_idx = 0

    # First part (before any h4)
    if parts[0].strip():
        text = html_to_text(parts[0])
        if '[TABLE]' in text:
            segments = text.split('[TABLE]')
            for i, seg in enumerate(segments):
                seg = seg.strip()
                if seg:
                    for line in seg.split('\n\n'):
                        line = line.strip()
                        if line:
                            doc.add_paragraph(line, style="Body Text")
                if i < len(segments) - 1 and table_idx < len(table_data):
                    headers, rows = table_data[table_idx]
                    add_table_to_doc(doc, headers, rows)
                    table_idx += 1
        else:
            for line in text.split('\n\n'):
                line = line.strip()
                if line:
                    doc.add_paragraph(line, style="Body Text")

    # Process h4 sections
    for i in range(1, len(parts), 2):
        heading_text = html_mod.unescape(re.sub(r'<[^>]+>', '', parts[i]).strip())
        doc.add_heading(heading_text, level=3)

        if i + 1 < len(parts):
            text = html_to_text(parts[i + 1])
            if '[TABLE]' in text:
                segments = text.split('[TABLE]')
                for j, seg in enumerate(segments):
                    seg = seg.strip()
                    if seg:
                        for line in seg.split('\n\n'):
                            line = line.strip()
                            if line:
                                doc.add_paragraph(line, style="Body Text")
                    if j < len(segments) - 1 and table_idx < len(table_data):
                        headers, rows = table_data[table_idx]
                        add_table_to_doc(doc, headers, rows)
                        table_idx += 1
            else:
                for line in text.split('\n\n'):
                    line = line.strip()
                    if line:
                        doc.add_paragraph(line, style="Body Text")

    # Add any remaining tables
    while table_idx < len(table_data):
        headers, rows = table_data[table_idx]
        add_table_to_doc(doc, headers, rows)
        table_idx += 1


def create_document():
    html = HTML_FILE.read_text(encoding="utf-8")
    sections = extract_section_titles(html)
    work_blocks = extract_your_work_sections(html)

    print(f"Found {len(sections)} sections: {[s[1] for s in sections]}")
    print(f"Found {len(work_blocks)} Your Work blocks")

    doc = Document()

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # -- Default font --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    body_style = doc.styles["Body Text"]
    body_style.font.name = "Calibri"
    body_style.font.size = Pt(11)
    body_style.paragraph_format.space_after = Pt(6)
    body_style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Calibri"
        hs.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ================================================================
    # COVER PAGE
    # ================================================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Assignment 1")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x48, 0x72, 0x65)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("GenAI-Powered RAG Assistant\nfor a DBA Gen AI Professor")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    doc.add_paragraph()

    details = [
        ("Student", "Sanjay Gupta"),
        ("Course", "DBA Gen AI: Pre-Trained Models (Course 8919)"),
        ("University", "Golden Gate University (GGU)"),
        ("Date", date.today().strftime("%B %d, %Y")),
        ("Job Selected", "Job #5: Secondary School Teacher → DBA Gen AI Professor"),
    ]

    for label, value in details:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}: ")
        run.bold = True
        run.font.size = Pt(12)
        run = p.add_run(value)
        run.font.size = Pt(12)

    doc.add_page_break()

    # ================================================================
    # CONTENT SECTIONS (1-8)
    # ================================================================
    # Map section numbers to their work blocks
    # Section 9 is checklist only — skip it
    content_sections = min(len(sections), 8)
    work_idx = 0

    for i in range(content_sections):
        sec_num, sec_title = sections[i]
        print(f"Processing Section {sec_num}: {sec_title}")

        doc.add_heading(f"Section {sec_num}: {sec_title}", level=1)

        if work_idx < len(work_blocks):
            block = work_blocks[work_idx]
            work_idx += 1

            # Special handling for Section 5 — insert architecture diagram
            if sec_num == "5":
                # Remove the figure tag from content (we'll add diagram separately)
                clean_block = re.sub(r'<figure[^>]*>.*?</figure>', '', block, flags=re.DOTALL)
                # Remove the "Architecture Diagram" h4 since we handle it
                clean_block = re.sub(r'<h4>Architecture Diagram</h4>', '', clean_block)
                process_section_content(doc, clean_block, sec_num)
                # Add diagram
                doc.add_heading("Architecture Diagram", level=3)
                try:
                    png_bytes = cairosvg.svg2png(
                        url=str(DIAGRAM_SVG),
                        output_width=1400
                    )
                    doc.add_picture(io.BytesIO(png_bytes), width=Inches(6.0))
                    last_para = doc.paragraphs[-1]
                    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap = doc.add_paragraph(
                        "Figure 1: 8-Stage RAG Pipeline Architecture",
                        style="Body Text"
                    )
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.runs[0].italic = True
                    cap.runs[0].font.size = Pt(9)
                except Exception as e:
                    doc.add_paragraph(f"[Diagram could not be embedded: {e}]")

            # Special handling for Section 8 — insert screenshots
            elif sec_num == "8":
                # Process content up to where screenshots go
                # Split at figure tags
                clean_block = re.sub(
                    r'<figure[^>]*>.*?</figure>',
                    '\n[SCREENSHOT]\n',
                    block,
                    flags=re.DOTALL
                )

                screenshot_keys = list(SCREENSHOTS.keys())
                ss_idx = 0

                segments = clean_block.split('[SCREENSHOT]')
                for j, seg in enumerate(segments):
                    seg = seg.strip()
                    if seg:
                        process_section_content(doc, seg, sec_num)

                    if j < len(segments) - 1 and ss_idx < len(screenshot_keys):
                        # Insert screenshot
                        fname = screenshot_keys[ss_idx]
                        caption = SCREENSHOTS[fname]
                        img_path = SCREENSHOT_DIR / fname
                        try:
                            doc.add_picture(str(img_path), width=Inches(5.5))
                            last_para = doc.paragraphs[-1]
                            last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cap = doc.add_paragraph(caption, style="Body Text")
                            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cap.runs[0].italic = True
                            cap.runs[0].font.size = Pt(9)
                        except Exception as e:
                            doc.add_paragraph(f"[Screenshot {fname} could not be embedded: {e}]")
                        ss_idx += 1
            else:
                process_section_content(doc, block, sec_num)

        # Page break between sections (except after last)
        if i < content_sections - 1:
            doc.add_page_break()

    # ================================================================
    # SAVE
    # ================================================================
    doc.save(str(OUTPUT))
    print(f"\nSaved: {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    create_document()
