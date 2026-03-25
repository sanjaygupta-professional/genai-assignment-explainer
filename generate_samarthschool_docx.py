#!/usr/bin/env python3.12
"""
Generate SamarthSchool Group Assignment DOCX from group-assignment-report.md.

Reads the markdown report and converts it to a formatted Word document with:
- DesignArena-styled cover page
- Embedded publication-quality diagrams (fig1-fig9)
- Teal-header tables, proper typography, styled headings
- Page header and footer with page numbers
- Inline citation rendering (superscript [N])
- Clickable hyperlinks in references
"""

import re
import io
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = Path("/home/opc/genai-assignment-explainer")
REPORT_MD = BASE / "group-assignment-report.md"
DIAGRAM_DIR = BASE / "images" / "diagrams" / "group-assignment"
OUTPUT = BASE / "outputs" / "SamarthSchool_Group_Assignment_v2.docx"
OUTPUT.parent.mkdir(parents=True, exist_ok=True)

# DesignArena palette
TEAL = RGBColor(0x48, 0x72, 0x65)
TEAL_DARK = RGBColor(0x35, 0x55, 0x4B)
DARK = RGBColor(0x29, 0x2C, 0x33)
GRAY = RGBColor(0x6B, 0x72, 0x80)
CREAM_HEX = "F7F6F5"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Map figure references to diagram files and captions
DIAGRAM_MAP = {
    "fig1": ("fig1-rag-pipeline", "Figure 1: SamarthSchool GraphRAG Pipeline — Dual retrieval with KG eligibility matching and vector-based semantic search"),
    "fig2": ("fig2-system-architecture", "Figure 2: System Architecture — End-to-end pipeline from data ingestion through Knowledge Graph and vector storage to multilingual response generation"),
    "fig3": ("fig3-data-flow", "Figure 3: Multilingual Query Pipeline — Language detection, cross-lingual retrieval, and response generation across Indian languages"),
    "fig4": ("fig4-future-architecture", "Figure 4: Knowledge Graph Schema — Entity-relationship structure for 50+ Indian disability welfare schemes"),
    "fig5": ("fig5-teacher-journey", "Figure 5: Human-in-the-Loop Workflow — Three-level design ensuring human oversight at KG construction, query review, and quality improvement"),
    "fig6": ("fig6-roadmap-gantt", "Figure 6: School Administrator Journey — From natural language query to personalized benefits action plan in under 5 minutes"),
    "fig7": ("fig7-roi-comparison", "Figure 7: Social Impact ROI Model — Three-year cost-benefit trajectory from 0.20x to 2.2x social ROI"),
    "fig8": ("fig8-benchmarking", "Figure 8: Competitive Landscape — How SamarthSchool addresses gaps left by existing solutions"),
    "fig9": ("fig9-social-enterprise", "Figure 9: Business and Funding Model — Revenue sources, adoption channels, and social impact trajectory"),
}

# Map section headings to diagrams for inline insertion
SECTION_DIAGRAM_INSERT = {
    "4.1 Why RAG": "fig1",
    "4.2 Architecture overview": "fig2",
    "4.4 Knowledge Graph schema": "fig4",
    "4.5 Multilingual pipeline": "fig3",
    "4.11 Human-in-the-loop": "fig5",
    "5.1 Phased development": "fig6",
    "5.2 Roadmap timeline": "fig7",  # reuse fig7 (roadmap gantt) here too
    "6.3 Social impact ROI": "fig7",
    "7.1 Existing solutions": "fig8",
    "7.3 Data and IP moat": "fig9",
}


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_hyperlink(paragraph, url, text, font_size=Pt(11), color=None):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>')
    run_elem = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr>'
        f'    <w:rStyle w:val="Hyperlink"/>'
        f'    <w:color w:val="{color or "0563C1"}"/>'
        f'    <w:u w:val="single"/>'
        f'    <w:sz w:val="{int(font_size.pt * 2)}"/>'
        f'  </w:rPr>'
        f'  <w:t xml:space="preserve">{text}</w:t>'
        f'</w:r>'
    )
    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def add_header_footer(doc):
    """Add page header text and footer with page number to all sections."""
    for section in doc.sections:
        # Different first page (cover page gets no header/footer)
        section.different_first_page_header_footer = True

        # Header for subsequent pages
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.text = ""
        run = hp.add_run("SamarthSchool - DBA Gen AI Group Assignment")
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        run.font.name = "Calibri"
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add bottom border to header
        pPr = hp._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="487265"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

        # Footer with page number
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        run.font.name = "Calibri"
        # Insert PAGE field
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fld_char_begin)
        instr_text = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run._r.append(instr_text)
        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run._r.append(fld_char_end)


def add_styled_table(doc, headers, rows):
    """Add a formatted table with teal header row."""
    ncols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        set_cell_shading(cell, "487265")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = WHITE

    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            if j < ncols:
                cell = table.cell(i + 1, j)
                cell.text = str(val).strip()
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
                if i % 2 == 1:
                    set_cell_shading(cell, CREAM_HEX)

    doc.add_paragraph()
    return table


def try_add_diagram(doc, fig_key, width=5.5):
    """Embed a diagram PNG by figure key."""
    if fig_key not in DIAGRAM_MAP:
        return False
    filename, caption = DIAGRAM_MAP[fig_key]
    img_path = DIAGRAM_DIR / f"{filename}.png"
    if not img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[{caption} — image not found]")
        run.italic = True
        run.font.color.rgb = GRAY
        return False
    try:
        doc.add_picture(str(img_path), width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
        return True
    except Exception as e:
        doc.add_paragraph(f"[{caption} — embed error: {e}]")
        return False


def add_cover_page(doc):
    """Add a styled cover page."""
    # Spacer
    for _ in range(6):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SamarthSchool")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = TEAL
    run.font.name = "Georgia"

    # Hindi
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("समर्थ स्कूल")
    run.font.size = Pt(20)
    run.font.color.rgb = GRAY
    run.font.name = "Georgia"

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AI-Powered Benefits Navigator for Children with Special Abilities")
    run.font.size = Pt(16)
    run.font.color.rgb = DARK

    doc.add_paragraph()

    # Horizontal rule
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 40)
    run.font.color.rgb = TEAL
    run.font.size = Pt(10)

    doc.add_paragraph()

    # Course info
    lines = [
        ("Group Assignment Report", Pt(14), True, DARK),
        ("Gen AI: Pre-Trained Models (Course 8919)", Pt(12), False, GRAY),
        ("GGU DBA Program via upGrad", Pt(12), False, GRAY),
        ("", Pt(8), False, GRAY),
        (f"Date: {date.today().strftime('%B %d, %Y')}", Pt(11), False, GRAY),
    ]
    for text, size, bold, color in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = size
        run.bold = bold
        run.font.color.rgb = color

    doc.add_page_break()


def parse_markdown_table(lines):
    """Parse markdown table lines into (headers, rows)."""
    if len(lines) < 2:
        return None, None
    headers = [c.strip() for c in lines[0].split("|")[1:-1]]
    # Skip separator line (line 1)
    rows = []
    for line in lines[2:]:
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if cells:
            rows.append(cells)
    return headers, rows


def process_inline_formatting(paragraph, text):
    """Add text to paragraph with bold/italic/citation inline formatting."""
    # Split on **bold**, *italic*, and [N] citation patterns
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|\[\d+\])', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(11)
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.size = Pt(11)
        elif re.match(r'^\[\d+\]$', part):
            # Render citation as superscript
            run = paragraph.add_run(part)
            run.font.size = Pt(8)
            run.font.superscript = True
            run.font.color.rgb = TEAL
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(11)


def convert_md_to_docx():
    """Main conversion: read markdown, produce formatted DOCX."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Styles
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.name = "Calibri"
        if level == 1:
            hs.font.color.rgb = TEAL
            hs.font.size = Pt(18)
        elif level == 2:
            hs.font.color.rgb = TEAL_DARK
            hs.font.size = Pt(14)
        else:
            hs.font.color.rgb = DARK
            hs.font.size = Pt(12)

    # Cover page
    add_cover_page(doc)

    # Header and footer (applied after cover page is created)
    add_header_footer(doc)

    # Read markdown
    md_text = REPORT_MD.read_text(encoding="utf-8")
    lines = md_text.split("\n")

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []
    diagrams_inserted = set()
    pending_diagram = None  # diagram to insert after next content paragraph
    in_references = False   # track when we're in the References section

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block — add as formatted text
                code_text = "\n".join(code_lines)
                if code_text.strip():
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                    run.font.color.rgb = DARK
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                # Insert pending diagram after code block too
                if pending_diagram:
                    doc.add_paragraph()
                    try_add_diagram(doc, pending_diagram, width=5.5)
                    doc.add_paragraph()
                    pending_diagram = None
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table handling
        if "|" in line and line.strip().startswith("|") and line.strip().endswith("|"):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # Table ended — render it
            headers, rows = parse_markdown_table(table_lines)
            if headers and rows:
                add_styled_table(doc, headers, rows)
            in_table = False
            table_lines = []
            # Don't increment i — reprocess current line
            continue

        # Skip title line (handled by cover page)
        if i == 0 and line.startswith("# "):
            i += 1
            continue

        # Horizontal rules
        if line.strip() == "---":
            i += 1
            continue

        # Headings
        if line.startswith("#### "):
            text = line[5:].strip()
            h = doc.add_heading(text, level=3)
            for run in h.runs:
                run.font.color.rgb = DARK
            i += 1
            continue

        if line.startswith("### "):
            text = line[4:].strip()
            h = doc.add_heading(text, level=3)
            for run in h.runs:
                run.font.color.rgb = DARK
            # Check if this section should have a diagram inserted
            for section_key, fig_key in SECTION_DIAGRAM_INSERT.items():
                if section_key in text and fig_key not in diagrams_inserted:
                    pending_diagram = fig_key
                    diagrams_inserted.add(fig_key)
                    break
            i += 1
            continue

        if line.startswith("## "):
            text = line[3:].strip()
            # Track references section for URL rendering
            in_references = "References" in text
            # Strip numbering prefix like "1. " for cleaner headings
            clean_text = re.sub(r"^\d+\.\s+", "", text)
            h = doc.add_heading(clean_text, level=1)
            for run in h.runs:
                run.font.color.rgb = TEAL
            i += 1
            continue

        # Bullet points
        if line.strip().startswith("- "):
            text = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            process_inline_formatting(p, text)
            i += 1
            continue

        # Numbered items (including APA references with URLs)
        m = re.match(r"^(\d+)\.\s+(.+)", line.strip())
        if m and not line.startswith("#"):
            text = m.group(2)
            # Check if this is an APA reference line with a URL at the end
            url_match = re.search(r'(https?://\S+)$', text)
            if url_match and in_references:
                url = url_match.group(1)
                pre_url = text[:url_match.start()].rstrip()
                p = doc.add_paragraph(style="List Number")
                process_inline_formatting(p, pre_url + " ")
                add_hyperlink(p, url, url, font_size=Pt(10))
            else:
                p = doc.add_paragraph(style="List Number")
                process_inline_formatting(p, text)
            i += 1
            continue

        # Empty lines
        if not line.strip():
            i += 1
            continue

        # Regular paragraphs
        text = line.strip()
        p = doc.add_paragraph()
        process_inline_formatting(p, text)

        # Insert pending diagram after the first content paragraph following the heading
        if pending_diagram:
            doc.add_paragraph()  # spacing
            try_add_diagram(doc, pending_diagram, width=5.5)
            doc.add_paragraph()  # spacing
            pending_diagram = None

        i += 1

    # Flush any remaining table
    if in_table and table_lines:
        headers, rows = parse_markdown_table(table_lines)
        if headers and rows:
            add_styled_table(doc, headers, rows)

    # Insert diagrams at the end as an appendix if not already embedded
    remaining_figs = [k for k in DIAGRAM_MAP if k not in diagrams_inserted]
    if remaining_figs:
        doc.add_page_break()
        h = doc.add_heading("Appendix: Architecture Diagrams", level=1)
        for run in h.runs:
            run.font.color.rgb = TEAL

        p = doc.add_paragraph()
        run = p.add_run(
            "The following publication-quality diagrams illustrate key aspects of the "
            "SamarthSchool architecture, data flow, roadmap, and business model."
        )
        run.font.size = Pt(11)
        run.font.color.rgb = GRAY
        doc.add_paragraph()

        for fig_key in DIAGRAM_MAP:
            try_add_diagram(doc, fig_key, width=5.8)
            doc.add_paragraph()

    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size / 1024 / 1024:.1f} MB")


if __name__ == "__main__":
    convert_md_to_docx()
